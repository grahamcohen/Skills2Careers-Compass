import json
import os
from datetime import datetime
import re

# Dependency Check
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("❌ Error: The 'python-docx' library is required.")
    print("👉 Please run: pip install python-docx")
    exit(1)

# Check for matplotlib
try:
    import matplotlib.pyplot as plt
    HAS_MATPLOTLIB = True
except ImportError:
    HAS_MATPLOTLIB = False
    print("⚠️ Warning: 'matplotlib' library not found. Charts will be skipped.")
    print("👉 To enable charts, run: pip install matplotlib")

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_FILE = os.path.join(BASE_DIR, "courses.json")
SECTOR_FILE = os.path.join(BASE_DIR, "sector_data.json")
# Resource Files for Ecosystem Analysis
RES_FILES = {
    "General": os.path.join(BASE_DIR, "resources_general.json"),
    "Digital": os.path.join(BASE_DIR, "resources_digital.json"),
    "Agriculture": os.path.join(BASE_DIR, "resources_agri.json"),
    "Energy": os.path.join(BASE_DIR, "resources_energy.json")
}
OUTPUT_FILE = os.path.join(BASE_DIR, "Skills_Landscape_Analysis.docx")
CHART_FILE_SECTOR = os.path.join(BASE_DIR, "sector_chart.png")
CHART_FILE_TYPE = os.path.join(BASE_DIR, "type_chart.png")
CHART_FILE_MAP = os.path.join(BASE_DIR, "map_chart.png")

def parse_duration(dur):
    if not dur:
        return None
    dur = dur.lower()
    match = re.search(r'[\d\.]+', dur)
    if not match:
        return None
    
    num = float(match.group(0))
    if 'year' in dur:
        return num * 12
    if 'month' in dur:
        return num
    if 'week' in dur:
        return num / 4.33
    if 'day' in dur:
        return num / 30
    if 'hour' in dur:
        return num / 730
    return None

def analyze_resources(sector_key):
    stats = {"Incubation": 0, "Mentorship": 0, "Funding": 0, "Jobs": 0}
    if os.path.exists(SECTOR_FILE):
        with open(SECTOR_FILE, 'r', encoding='utf-8') as f:
            sector_demand = json.load(f)
            sector_info = sector_demand.get(sector_key, {})
            resources = sector_info.get("resources", {})
            
            communities = resources.get("communities", [])
            stats["Mentorship"] = len(communities)
            
            jobs = resources.get("jobs", [])
            stats["Jobs"] = len(jobs)
            
            entrepreneurship = resources.get("entrepreneurship", {})
            stats["Incubation"] = len(entrepreneurship.get("incubators", []))
            stats["Funding"] = len(entrepreneurship.get("funding", []))
            
    return stats

def create_document():
    # 1. Load Data
    data = []
    course_count = 0
    sector_counts = {}
    if os.path.exists(DATA_FILE):
        with open(DATA_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
            course_count = len(data)
            for c in data:
                s = c.get('sector', 'Other')
                if s == 'agri': label = 'Agriculture'
                elif s == 'energy': label = 'Energy'
                elif s == 'digital': label = 'Digital'
                elif s == 'all': label = 'Multi-Sector'
                else: label = s.capitalize()
                sector_counts[label] = sector_counts.get(label, 0) + 1    
    else:
        print(f"⚠️ Warning: {DATA_FILE} not found. Metadata will be incomplete.")

    # Load Sector Data (Demand)
    sector_demand = {}
    if os.path.exists(SECTOR_FILE):
        with open(SECTOR_FILE, 'r', encoding='utf-8') as f:
            sector_demand = json.load(f)

    # --- Analysis ---
    # Geographic Analysis
    country_counts = {}
    for course in data:
        country = course.get('country', 'Unknown')
        country_label = country.capitalize() if country != 'all' else 'Regional/Global'
        country_counts[country_label] = country_counts.get(country_label, 0) + 1

    # Course Type Analysis
    type_categories = {
        'Degree/Diploma': ['degree', 'diploma'],
        'Certificate/License': ['certificate', 'license', 'certification', 'professional certificate'],
        'Bootcamp/Practical': ['bootcamp', 'initiative', 'fellowship'],
        'Short Course': ['short course', 'workshop', 'micromodule', 'micro-credential'],
        'Platform/Community': ['platform', 'community', 'academy'],
        'Incubator/Hub': ['incubator', 'hub', 'lab'],
        'TVET': ['tvet', 'polytechnic']
    }
    course_type_counts = {cat: 0 for cat in type_categories}
    course_type_counts['Other'] = 0
    for c in data:
        course_type_lower = c.get('type', '').lower()
        found = False
        for cat, keywords in type_categories.items():
            if any(kw in course_type_lower for kw in keywords):
                course_type_counts[cat] += 1
                found = True
                break
        if not found:
            course_type_counts['Other'] += 1

    # Duration Analysis
    duration_buckets = {
        "< 1 Month": 0, "1-3 Months": 0, "3-6 Months": 0,
        "6-12 Months": 0, "1-2 Years": 0, "2+ Years": 0, "Self-paced/Unknown": 0
    }
    for c in data:
        months = parse_duration(c.get('duration'))
        if months is None:
            duration_buckets["Self-paced/Unknown"] += 1
        elif months < 1:
            duration_buckets["< 1 Month"] += 1
        elif months <= 3:
            duration_buckets["1-3 Months"] += 1
        elif months <= 6:
            duration_buckets["3-6 Months"] += 1
        elif months <= 12:
            duration_buckets["6-12 Months"] += 1
        elif months <= 24:
            duration_buckets["1-2 Years"] += 1
        else:
            duration_buckets["2+ Years"] += 1

    # --- Document Generation ---
    doc = Document()
    title = doc.add_heading('Skills Training Landscape Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Generated for Skills2Careers Compass | Dataset: {course_count} Courses\n")
    run.italic = True
    run.font.size = Pt(9)
    meta.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}")

    add_heading(doc, "Executive Summary", 1)
    doc.add_paragraph(
        "The analysis of the Skills2Careers Compass training catalog reveals a landscape dominated by Digital skills, largely delivered through scalable online platforms. "
        "This creates high accessibility but raises questions about local labor market relevance and practical competency. In contrast, the Agriculture and Energy sectors, "
        "while anchored by traditional university degrees, exhibit a significant 'missing middle' of agile, short-term vocational training needed to meet immediate industry demands. "
        "Geographically, Kenya serves as a regional hub with a diverse provider ecosystem, while other EAC nations show greater reliance on public institutions and global online content. "
        "A key policy opportunity lies in creating frameworks that recognize and stack micro-credentials from various providers, bridging the gap between formal education and job-ready skills."
    )

    add_heading(doc, "1. Overall Training Landscape", 1)
    add_heading(doc, "1.1. Distribution by Sector", 2)
    if HAS_MATPLOTLIB and sector_counts:
        if generate_chart(sector_counts, "Course Distribution by Sector", CHART_FILE_SECTOR):
            doc.add_picture(CHART_FILE_SECTOR, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if os.path.exists(CHART_FILE_SECTOR):
                os.remove(CHART_FILE_SECTOR)
    doc.add_paragraph("The training catalog is heavily concentrated in the Digital Economy, reflecting global trends and the scalability of online tech education. Multi-Sector courses, often foundational digital skills or entrepreneurship, also represent a significant portion. Agriculture and Energy have a smaller, more specialized footprint.")

    add_heading(doc, "1.2. Distribution by Geography", 2)
    
    if HAS_MATPLOTLIB and country_counts:
        if generate_map(country_counts, CHART_FILE_MAP):
            doc.add_picture(CHART_FILE_MAP, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if os.path.exists(CHART_FILE_MAP):
                os.remove(CHART_FILE_MAP)

    add_table(doc, ["Country", "Number of Courses"], sorted(country_counts.items(), key=lambda item: item[1], reverse=True))
    doc.add_paragraph("A large number of courses are 'Regional/Global', primarily online offerings from international platforms. Among specific countries, Kenya has the most diverse and numerous training options, positioning it as a regional hub. Other nations show a developing but smaller local training ecosystem.")

    add_heading(doc, "1.3. Distribution by Course Type", 2)
    add_table(doc, ["Course Type", "Count"], sorted(course_type_counts.items(), key=lambda item: item[1], reverse=True))
    doc.add_paragraph("Short Courses and Certificates dominate, indicating a market preference for rapid, targeted upskilling. The high number of 'Platforms' reflects the inclusion of large MOOCs. Traditional Degrees/Diplomas form a stable but smaller part of the agile training landscape.")

    add_heading(doc, "1.4. Distribution by Duration", 2)
    add_table(doc, ["Duration", "Count"], duration_buckets.items())
    doc.add_paragraph("The landscape is polarized between long-term academic programs (2+ years) and a high volume of short-term (<3 months) and self-paced courses. This highlights a 'missing middle' of medium-term (6-12 month) practical diplomas or apprenticeships that are crucial for deep vocational skill-building.")

    add_heading(doc, "2. Deep Dive: Sector-Specific Supply vs. Demand", 1)
    
    if sector_demand and data:
        for sector_key, sector_info in sector_demand.items():
            if not isinstance(sector_info, dict): continue
            
            sector_name = sector_info.get('meta', {}).get('name', sector_key.capitalize())
            add_heading(doc, f"2.1. {sector_name} Ecosystem", 2)
            
            growth = sector_info.get('growth', {})
            p = doc.add_paragraph()
            p.add_run("Market Signal: ").bold = True
            p.add_run(f"Job trend is {growth.get('jobTrend', 'N/A')} with investment rated as {growth.get('investment', 'N/A')}. Overall skills demand is '{growth.get('skillsDemand', 'N/A')}'.")
            
            sector_courses = [c for c in data if c.get('sector') == sector_key]
            course_count = len(sector_courses)
            
            hot_skills = [s['skill'] for s in sector_info.get('skills', []) if s.get('isHot')]
            covered_hot_skills = []
            missing_hot_skills = []
            
            for skill in hot_skills:
                match = any(skill.lower() in (s.lower() for s in c.get('skills', [])) for c in sector_courses)
                if match:
                    covered_hot_skills.append(skill)
                else:
                    missing_hot_skills.append(skill)
            
            p = doc.add_paragraph()
            p.add_run("Training Supply Analysis: ").bold = True
            p.add_run(f"{course_count} sector-specific courses were identified. ")
            
            if missing_hot_skills:
                p.add_run(f"A potential supply gap exists for high-demand skills like: {', '.join(missing_hot_skills)}. ")
            else:
                p.add_run("Training supply appears well-aligned with high-demand skills. ")

            # Support Ecosystem
            p = doc.add_paragraph()
            p.add_run("Support Ecosystem: ").bold = True
            support_stats = analyze_resources(sector_key)
            p.add_run(f"Identified {support_stats['Incubation']} incubation/hub resources, {support_stats['Funding']} funding opportunities, and {support_stats['Mentorship']} mentorship/community networks.")

    add_heading(doc, "3. Provider & Modality Analysis", 1)
    add_heading(doc, "3.1. Provider Typology", 2)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Sector'
    hdr[1].text = 'Incubation/Hubs'
    hdr[2].text = 'Mentorship/Networks'
    hdr[3].text = 'Funding/Grants'
    hdr[4].text = 'Jobs/Placement'
    
    for sector_key in ["digital", "agri", "energy"]:
        sector_name = "Digital" if sector_key == "digital" else "Agriculture" if sector_key == "agri" else "Energy"
        stats = analyze_resources(sector_key)
        row = table.add_row().cells
        row[0].text = sector_name
        row[1].text = str(stats['Incubation'])
        row[2].text = str(stats['Mentorship'])
        row[3].text = str(stats['Funding'])
        row[4].text = str(stats['Jobs'])

    add_heading(doc, "3.2. Modality & Accessibility", 2)
    doc.add_paragraph("• Online Dominance: The majority of 'new economy' skills (Data, AI, Cloud) are delivered Online, offering scalability but posing challenges for learners with limited connectivity or need for hands-on practice.")
    doc.add_paragraph("• The 'Hybrid' Necessity: Successful employment outcomes often correlate with Hybrid models that combine online theory with in-person mentorship and practical projects (e.g., Moringa School, Refactory).")
    doc.add_paragraph("• Language Gap: The catalog is overwhelmingly English. French and Kiswahili technical training options are sparse, creating a barrier to entry for many learners in the EAC.")

    add_heading(doc, "4. Conclusion & Recommendations", 1)
    doc.add_paragraph(
        "The data suggests a two-speed skills ecosystem: a fast-moving, globalized digital track and a slower, more traditional track for Agriculture and Energy. To foster inclusive growth, policy should focus on bridging this divide."
    )
    add_bullet(doc, "Recommendation 1", "Establish a regional micro-credential framework. Create a system where learners can 'stack' verified skills from different providers (universities, bootcamps, online platforms) into a recognized qualification. This increases flexibility and validates informal learning.")
    add_bullet(doc, "Recommendation 2", "Incentivize 'Missing Middle' TVET. Use public-private partnerships to fund the development of short-to-medium term (6-12 month) practical diplomas in Agriculture and Energy, focusing on skills gaps like IoT maintenance, drone operations, and mini-grid management.")
    add_bullet(doc, "Recommendation 3", "Promote Localized Hybrid Learning. Support the creation of physical 'learning hubs' in secondary cities and rural areas where students can access reliable internet, receive in-person mentorship, and collaborate on projects while taking global online courses.")

    doc.save(OUTPUT_FILE)
    print(f"✅ Report generated successfully: {OUTPUT_FILE}")

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    if level == 1:
        h.style.font.color.rgb = RGBColor(0, 51, 102) # Dark Blue
    elif level == 2:
        h.style.font.color.rgb = RGBColor(51, 102, 153) # Lighter Blue

def add_bullet(doc, label, text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f"{label}: ").bold = True
    p.add_run(text)

def add_row(table, c1, c2, c3):
    row = table.add_row().cells
    row[0].text = c1
    row[1].text = c2
    row[2].text = c3

def generate_chart(counts, title, filename):
    try:
        sectors = list(counts.keys())
        values = list(counts.values())
        
        # Sort by value descending
        sorted_pairs = sorted(zip(sectors, values), key=lambda x: x[1], reverse=True)
        sectors = [x[0] for x in sorted_pairs]
        values = [x[1] for x in sorted_pairs]

        plt.figure(figsize=(6, 4))
        colors = ['#4f46e5' if 'Digital' in s else '#10b981' if 'Agri' in s else '#f59e0b' if 'Energy' in s else '#64748b' for s in sectors] # Indigo, Green, Orange, Slate
        
        bars = plt.bar(sectors, values, color=colors)
        plt.title(title, fontsize=12, fontweight='bold')
        plt.ylabel('Number of Courses')
        plt.grid(axis='y', linestyle='--', alpha=0.3)
        
        plt.tight_layout()
        plt.savefig(filename, dpi=300)
        plt.close()
        return True
    except Exception as e:
        print(f"⚠️ Failed to generate chart: {e}")
        return False

def generate_map(country_counts, filename):
    try:
        import geopandas as gpd
    except ImportError:
        print("⚠️ Warning: 'geopandas' library not found. Map will be skipped.")
        print("👉 To enable map charts, run: pip install geopandas")
        return False
        
    try:
        # Load built-in world map
        world = gpd.read_file(gpd.datasets.get_path('naturalearth_lowres'))
        
        # Filter for relevant East African nations
        eac_countries = ['Kenya', 'Uganda', 'United Republic of Tanzania', 'Rwanda', 'Burundi', 'South Sudan', 'Dem. Rep. Congo', 'Somalia']
        eac = world[world.name.isin(eac_countries)].copy()
        
        # Map country names to match dataset
        mapped_counts = {}
        for k, v in country_counts.items():
            if k == 'Tanzania': mapped_counts['United Republic of Tanzania'] = v
            elif k.lower() in ['dr congo', 'drc']: mapped_counts['Dem. Rep. Congo'] = v
            else: mapped_counts[k] = v
            
        eac['Course_Count'] = eac['name'].map(mapped_counts).fillna(0)
        
        fig, ax = plt.subplots(1, 1, figsize=(6, 5))
        eac.plot(column='Course_Count', ax=ax, legend=True,
                 legend_kwds={'label': "Number of Courses", 'orientation': "horizontal"},
                 cmap='YlGnBu', edgecolor='black', missing_kwds={'color': 'lightgrey'})
        
        plt.title("Geographic Distribution of Courses", fontsize=12, fontweight='bold')
        ax.set_axis_off()
        plt.tight_layout()
        plt.savefig(filename, dpi=300)
        plt.close()
        return True
    except Exception as e:
        print(f"⚠️ Failed to generate map: {e}")
        return False

if __name__ == "__main__":
    create_document()